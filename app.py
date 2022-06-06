# 弱小和无知不是生存的障碍，傲慢才是!
# coding=utf-8
import json
import os
import base64
import shutil
import datetime
import docx
from flask import Flask, request
from ocr import api_call
from filereader import FILEREADER
from win32com import client
import pythoncom


app = Flask(__name__)

file_reader_engine = 'http://117.161.12.146:8516/service/'
ocr_server = "http://117.161.12.146:8513"
scene = "chinese_print"
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
result = []
res = None


@app.route('/')
def index():
    return "Hello GuoDianTou!"


@app.route("/GetGdt", methods=["GET"])
def GetGdt():
    invalid_data_error0 = {'status': 'success', 'message': 'ok', 'data': None}
    return json.dumps(invalid_data_error0, indent=2, ensure_ascii=False, sort_keys=True)


def WriteFileGrade(filepath):
    if os.path.getsize(BASE_DIR + "\\" + "fileGrade.log") > 1073741824*2:
        os.remove(BASE_DIR + "\\" + "fileGrade.log")
    with open("fileGrade.log", "a+", encoding='UTF-8') as f:
        f.write(datetime.datetime.now().strftime('%b-%d-%Y-%H-%M-%S')+"，")
        f.write("文件路径（附带名字）:"+filepath+"，")
        f.write("所有敏感词+等级:"+str(result)+"，")
        f.write("最高等级:"+str(res))
        f.write('\n')


def function(date):
    return date['敏感等级']


def filter_value(msg):
    with open("filter.json", encoding='UTF-8') as f:
        content = json.load(f)
    flag = 0
    for grade in content:
        for val in content[grade]:
            if val in msg:
                flag = 1
                print(flag)
                result.append({"关键字": val, "敏感等级": grade, "数量": msg.count(val)})
                result.sort(key=function, reverse=True)
                for re in result:
                    global res
                    res = re
                    print(res)
                    invalid_data_error0 = {'status': 'success', 'message': 'ok', 'data': {"result": res}}
                    return json.dumps(invalid_data_error0, indent=2, ensure_ascii=False, sort_keys=True)
    if flag == 0:
        print("Error: No sensitive characters")
        invalid_data_error0 = {'status': 'success', 'message': 'ok',
                               'data': {"result": "Error: No sensitive characters"}}
        return json.dumps(invalid_data_error0, indent=2, ensure_ascii=False, sort_keys=True)


def Identification(file_path):
    file = open(file_path, "r", encoding="GB2312", errors='ignore')
    content = file.read()
    file.close()
    msg = "".join(content).replace("\n", "").replace(" ", "").replace("（","(").replace("）", ")")
    print(msg)
    result.clear()
    WriteFileGrade(file_path)
    return filter_value(msg)


def Docx_Ident(file_path):
    content = ""
    if os.path.splitext(file_path)[-1] == ".docx":
        document = docx.Document(file_path)
        for paragraph in document.paragraphs:
            content += paragraph.text
        msg = ''.join(content).replace(" ","").replace("\t","").replace("\n","")
        result.clear()
        WriteFileGrade(file_path)
        print(msg)
        return filter_value(msg)
    else:
        pythoncom.CoInitialize()
        word = client.Dispatch('Word.Application')
        word.Visible = 0  # 后台运行,不显示
        word.DisplayAlerts = 0  # 不警告
        doc = word.Documents.Open(file_path)
        for para in doc.paragraphs:
            print(para.Range.Text)
        doc.SaveAs('D:PythonFiles/4paradigm/gdt_flask/file/test.txt', 2)
        doc.Close()
        word.Quit()
        pythoncom.CoUninitialize()
        with open('D:/PythonFiles/4paradigm/gdt_flask/file/test.txt', 'r', encoding="gbk") as f:
            content = f.read()
            msg = "".join(content).replace("\n", "").replace(" ", "").replace("（", "(").replace("）", ")")
            print(msg)
            result.clear()
            WriteFileGrade(file_path)
            return filter_value(msg)


def Ocr_Ident(file_path, scene):
    content = api_call(file_path, scene)
    data = content['data']['json']['general_ocr_res']['texts']
    msg = ""
    for da in data:
        msg += da
    result.clear()
    WriteFileGrade(file_path)
    return filter_value(msg)


def Filereader_Ident(file_reader_engine, file_path):
    filereader = FILEREADER(file_reader_engine, r"%s" % file_path)
    filereader.f_upload()
    filereader.f_inquire()
    content = filereader.f_download()
    msg = "".join(content).replace("\n", "")
    result.clear()
    WriteFileGrade(file_path)
    return filter_value(msg)


def ClassifyModel(new_file):
    print(new_file)
    if os.path.splitext(new_file)[-1] == ".pdf":
        return Filereader_Ident(file_reader_engine, new_file)
    elif os.path.splitext(new_file)[-1] in ['.png', '.jpg', '.jpeg']:
        return Ocr_Ident(new_file, scene)
    elif os.path.splitext(new_file)[-1] in ['.txt']:
        return Identification(new_file)
    elif os.path.splitext(new_file)[-1] in ['.doc', '.docx']:
        return Docx_Ident(new_file)
    else:
        invalid_data_error0 = {'status': 'fail', 'message': 'File suffix error', 'data': {"result": "Error: NotFound suffix"}}
        return json.dumps(invalid_data_error0, indent=2, ensure_ascii=False, sort_keys=True)


def WriteFile(suffix, file):
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
    if len(os.listdir(BASE_DIR + "\\file")) != 0:
        shutil.rmtree(BASE_DIR + "\\file")
        os.mkdir(BASE_DIR + "\\file")
    with open(BASE_DIR + "\\file\\recognize_gdt" + suffix, "wb") as f:
        f.write(base64.b64decode(file))
    return ClassifyModel(BASE_DIR + "\\file\\recognize_gdt" + suffix)


@app.route("/PostGdt", methods=["POST", 'GET'])
def PostGdt():
    if not request.data:
        invalid_data_error0 = {'status': 'fail', 'message': 'bad_request_method', 'data': None}
        return json.dumps(invalid_data_error0, indent=2, ensure_ascii=False, sort_keys=True)
    req_input = json.loads(request.data.decode('utf-8'))
    # if req_input['suffix'] not in ['.txt', '.doc', '.docx', '.pdf', '.jpg', '.jpeg', '.png']:
    #     invalid_data_error0 = {'status': 'fail', 'message': 'File suffix error', 'data': None}
    #     # return json.dumps(invalid_data_error0, indent=2, ensure_ascii=False, sort_keys=True)
    # else:
    return WriteFile(req_input['suffix'], req_input['file'])


if __name__ == '__main__':
    app.run(debug=True)